from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    KeepInFrame,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "cv_designed.docx"
PDF_PATH = ROOT / "cv_designed.pdf"

ACCENT = "#18324b"
ACCENT_MUTED = "#5d7288"
SIDEBAR_BG = "#eef3f7"
TEXT = "#172433"
SUBTLE = "#546170"
WHITE = "#ffffff"


RESUME = {
    "name": "Loay Eldin Mohamed Mahmoud",
    "title": "Full-Stack Developer | Robotics Engineer | Technical Educator",
    "contact": [
        "Alexandria, Egypt",
        "Age 19",
        "loayeldin77@gmail.com",
        "+20 114 499 922",
        "github.com/loayden",
        "linkedin.com/in/loayeldin77",
    ],
    "summary": (
        "Full-stack developer and robotics engineer based in Alexandria, Egypt, "
        "with a track record across live web products, robotics innovation, and "
        "technical mentoring. Builds polished digital experiences with Next.js, "
        "React, and TypeScript, while also developing robotics and AI concepts "
        "using Arduino, Raspberry Pi, Python, sensors, and computer vision. "
        "Represented Egypt in international competitions, earned 25+ awards and "
        "podium finishes, and mentored 100+ students through robotics and programming workshops."
    ),
    "skills": [
        (
            "Frontend",
            "Next.js, React, TypeScript, Tailwind CSS, responsive systems, "
            "component architecture, visual hierarchy",
        ),
        (
            "Backend",
            "Node.js, REST APIs, App Router, server-side logic, form workflows, "
            "contact systems",
        ),
        (
            "Robotics & AI",
            "Arduino, Raspberry Pi, Python, C++, OpenCV, sensors, ChatGPT API, "
            "speech systems",
        ),
        (
            "Product & Mentoring",
            "UI structure, technical communication, workshops, mentorship, "
            "competition preparation",
        ),
    ],
    "experience": [
        {
            "role": "Independent Full-Stack Developer & Product Builder",
            "organization": "Self-directed and freelance work",
            "period": "2023 - Present",
            "bullets": [
                "Designed and shipped public-facing websites and portfolio systems that turn rough ideas into clear, premium, and conversion-focused digital experiences.",
                "Owned frontend execution across structure, responsive behavior, visual hierarchy, and product communication using Next.js, React, and TypeScript.",
                "Delivered live web work including the VibeUp Event Platform and Aurelien ETA Fashion Website, with an emphasis on credibility, usability, and polished presentation.",
            ],
        },
        {
            "role": "Robotics Engineer, Prototype Builder & International Competitor",
            "organization": "Competition and prototype work",
            "period": "2018 - Present",
            "bullets": [
                "Designed robotics and AI concepts that address education, home care, accessibility, and assistive technology challenges.",
                "Represented Egypt internationally and achieved standout results including 1st place at MRC Greece and 2nd place worldwide at Robo Tourney USA.",
                "Built with Raspberry Pi, Arduino, Python, C++, sensors, and computer vision in high-pressure engineering, prototyping, and presentation environments.",
            ],
        },
        {
            "role": "Robotics Mentor & Technical Educator",
            "organization": "Workshops and student coaching",
            "period": "2024 - Present",
            "bullets": [
                "Mentored 100+ students in robotics, programming, and competition preparation through practical coaching, workshops, and project guidance.",
                "Turned difficult technical topics into clear, student-friendly lessons that improved understanding, confidence, and hands-on execution.",
            ],
        },
    ],
    "projects": [
        {
            "title": "VibeUp Event Platform",
            "meta": "Live website | Next.js, React",
            "bullets": [
                "Built a live event-platform website with stronger section hierarchy, clearer pricing communication, and a more premium product presentation.",
                "Live: https://vibeup-event.vercel.app",
            ],
        },
        {
            "title": "Aurelien ETA Fashion Website",
            "meta": "Live website | Next.js, React",
            "bullets": [
                "Created a premium fashion brand experience with editorial storytelling, refined visual hierarchy, and responsive layout execution.",
                "Live: https://aurelien-eta.vercel.app",
            ],
        },
        {
            "title": "AI Teacher Robot",
            "meta": "Robotics and EdTech concept | Raspberry Pi, Python, ChatGPT API, OpenCV",
            "bullets": [
                "Designed an interactive teaching robot that listens to students, answers questions by voice, and explains lessons step by step.",
                "Focused the concept on accessible learning for students who are sick, disabled, or underserved by traditional classroom methods.",
            ],
        },
        {
            "title": "VR Live Education Platform",
            "meta": "VR and EdTech concept | VR headset, online app, virtual classroom system",
            "bullets": [
                "Planned a VR-based learning platform that helps students learn from home, revisit missed lessons, and stay more focused through immersive interaction.",
                "Positioned the system as a more accessible and engaging option for students with illness, disabilities, or different learning needs.",
            ],
        },
        {
            "title": "HomeCare Assist Robot",
            "meta": "Healthcare robotics concept | Sensors, monitoring, patient support",
            "bullets": [
                "Designed a home-assistance robot for patients, seniors, and people with disabilities with cleaning, medicine reminders, food delivery, and safety alerts.",
                "Combined health monitoring, remote communication, and daily support features into one practical home-care concept.",
            ],
        },
    ],
    "education": [
        {
            "school": "Saxony Egypt University",
            "degree": "Automotive Engineering",
            "location": "Alexandria, Egypt",
        }
    ],
    "competitions": [
        {
            "title": "MRC Greece",
            "detail": "1st Place | Greece | 2024",
        },
        {
            "title": "Robo Tourney USA",
            "detail": "2nd Place Worldwide | California, USA | 2024",
        },
        {
            "title": "SeaPerch World Championship",
            "detail": "4th Place Worldwide, Best Maneuver, Best Team Spirit | 2022",
        },
        {
            "title": "Young Innovators Africa",
            "detail": "1st Place in Africa | 2022",
        },
        {
            "title": "Robot Challenge",
            "detail": "Represented Egypt in Beijing world championship | 2019, 2021, 2022",
        },
    ],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders_none(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "nil")


def add_text(paragraph, text, size, bold=False, color=TEXT, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    return run


def add_heading(container, text, sidebar=False):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if sidebar else 8)
    p.paragraph_format.space_after = Pt(3)
    add_text(p, text.upper(), 9.5, bold=True, color=ACCENT)


def build_docx():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    header = document.add_table(rows=1, cols=1)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    set_table_borders_none(header)
    header_cell = header.cell(0, 0)
    set_cell_shading(header_cell, ACCENT.replace("#", ""))
    set_cell_margins(header_cell, 170, 180, 170, 180)

    p = header_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_text(p, RESUME["name"], 22, bold=True, color=WHITE)

    p = header_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_text(p, RESUME["title"], 11, bold=True, color="E9F0F7")

    p = header_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_text(p, " | ".join(RESUME["contact"]), 9.2, color="E6EEF6")

    document.add_paragraph()

    body = document.add_table(rows=1, cols=2)
    body.alignment = WD_TABLE_ALIGNMENT.CENTER
    body.autofit = False
    set_table_borders_none(body)

    sidebar = body.cell(0, 0)
    main = body.cell(0, 1)

    sidebar.width = Inches(2.1)
    main.width = Inches(4.9)
    sidebar.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    main.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_shading(sidebar, SIDEBAR_BG.replace("#", ""))
    set_cell_margins(sidebar, 160, 170, 160, 170)
    set_cell_margins(main, 120, 180, 120, 180)

    add_heading(sidebar, "Skills", sidebar=True)
    for title, details in RESUME["skills"]:
        p = sidebar.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_text(p, title, 10, bold=True, color=TEXT)
        p = sidebar.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        add_text(p, details, 9.4, color=SUBTLE)

    add_heading(sidebar, "Education", sidebar=True)
    for edu in RESUME["education"]:
        p = sidebar.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_text(p, edu["school"], 10.2, bold=True, color=TEXT)
        p = sidebar.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_text(p, edu["degree"], 9.4, color=SUBTLE)
        p = sidebar.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        add_text(p, edu["location"], 9.2, color=ACCENT_MUTED)

    add_heading(sidebar, "Competitions", sidebar=True)
    for competition in RESUME["competitions"]:
        p = sidebar.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_text(p, competition["title"], 9.7, bold=True, color=TEXT)
        p = sidebar.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_text(p, competition["detail"], 8.9, color=ACCENT_MUTED)

    add_heading(main, "Summary")
    p = main.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_text(p, RESUME["summary"], 10, color=TEXT)

    add_heading(main, "Experience")
    for item in RESUME["experience"]:
        p = main.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_text(p, item["role"], 10.5, bold=True, color=TEXT)

        p = main.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_text(p, f'{item["organization"]} | {item["period"]}', 9.2, color=ACCENT_MUTED)

        for bullet in item["bullets"]:
            p = main.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            add_text(p, bullet, 9.4, color=TEXT)

    add_heading(main, "Projects")
    for project in RESUME["projects"]:
        p = main.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_text(p, project["title"], 10.4, bold=True, color=TEXT)

        p = main.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_text(p, project["meta"], 9.1, color=ACCENT_MUTED)

        for bullet in project["bullets"]:
            p = main.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            add_text(p, bullet, 9.3, color=TEXT)

    document.core_properties.author = RESUME["name"]
    document.core_properties.title = f'{RESUME["name"]} CV'
    document.save(DOCX_PATH)


def make_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="HeaderName",
            fontName="Helvetica-Bold",
            fontSize=21,
            leading=24,
            textColor=colors.HexColor(WHITE),
            alignment=TA_CENTER,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HeaderTitle",
            fontName="Helvetica-Bold",
            fontSize=10.4,
            leading=13,
            textColor=colors.HexColor("#EAF1F8"),
            alignment=TA_CENTER,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HeaderMeta",
            fontName="Helvetica",
            fontSize=8.4,
            leading=10.5,
            textColor=colors.HexColor("#E6EEF6"),
            alignment=TA_CENTER,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionHeading",
            fontName="Helvetica-Bold",
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor(ACCENT),
            spaceBefore=8,
            spaceAfter=5,
            allCaps=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            fontName="Helvetica",
            fontSize=8.8,
            leading=12,
            textColor=colors.HexColor(TEXT),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Meta",
            fontName="Helvetica",
            fontSize=8.1,
            leading=10,
            textColor=colors.HexColor(ACCENT_MUTED),
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Role",
            fontName="Helvetica-Bold",
            fontSize=9.2,
            leading=11.4,
            textColor=colors.HexColor(TEXT),
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SidebarLabel",
            fontName="Helvetica-Bold",
            fontSize=8.8,
            leading=10.8,
            textColor=colors.HexColor(TEXT),
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SidebarBody",
            fontName="Helvetica",
            fontSize=8.3,
            leading=10.8,
            textColor=colors.HexColor(SUBTLE),
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeBullet",
            fontName="Helvetica",
            fontSize=8.6,
            leading=11,
            textColor=colors.HexColor(TEXT),
            leftIndent=10,
            firstLineIndent=-8,
            spaceAfter=2,
        )
    )
    return styles


def bullet_paragraph(text, styles):
    return Paragraph(f"&bull; {text}", styles["ResumeBullet"])


def build_pdf():
    styles = make_pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=1.15 * cm,
        rightMargin=1.15 * cm,
        topMargin=1.15 * cm,
        bottomMargin=1.1 * cm,
        title=f'{RESUME["name"]} CV',
        author=RESUME["name"],
    )

    header_content = [
        Paragraph(RESUME["name"], styles["HeaderName"]),
        Paragraph(RESUME["title"], styles["HeaderTitle"]),
        Paragraph(" | ".join(RESUME["contact"]), styles["HeaderMeta"]),
    ]

    header = Table([[header_content]], colWidths=[doc.width])
    header.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(ACCENT)),
                ("BOX", (0, 0), (-1, -1), 0, colors.HexColor(ACCENT)),
                ("LEFTPADDING", (0, 0), (-1, -1), 18),
                ("RIGHTPADDING", (0, 0), (-1, -1), 18),
                ("TOPPADDING", (0, 0), (-1, -1), 14),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 13),
            ]
        )
    )

    header_width, header_height = header.wrap(doc.width, doc.topMargin)

    sidebar_story = [Paragraph("Skills", styles["SectionHeading"])]
    for title, details in RESUME["skills"]:
        sidebar_story.append(Paragraph(title, styles["SidebarLabel"]))
        sidebar_story.append(Paragraph(details, styles["SidebarBody"]))

    sidebar_story.append(Paragraph("Education", styles["SectionHeading"]))
    for edu in RESUME["education"]:
        sidebar_story.append(Paragraph(edu["school"], styles["SidebarLabel"]))
        sidebar_story.append(Paragraph(edu["degree"], styles["SidebarBody"]))
        sidebar_story.append(Paragraph(edu["location"], styles["Meta"]))

    sidebar_story.append(Paragraph("Competitions", styles["SectionHeading"]))
    for competition in RESUME["competitions"]:
        sidebar_story.append(Paragraph(competition["title"], styles["SidebarLabel"]))
        sidebar_story.append(Paragraph(competition["detail"], styles["Meta"]))

    main_story = [
        Paragraph("Summary", styles["SectionHeading"]),
        Paragraph(RESUME["summary"], styles["Body"]),
        Paragraph("Experience", styles["SectionHeading"]),
    ]

    for item in RESUME["experience"]:
        main_story.append(Paragraph(item["role"], styles["Role"]))
        main_story.append(
            Paragraph(f'{item["organization"]} | {item["period"]}', styles["Meta"])
        )
        for bullet in item["bullets"]:
            main_story.append(bullet_paragraph(bullet, styles))

    main_story.append(Paragraph("Projects", styles["SectionHeading"]))
    for project in RESUME["projects"]:
        main_story.append(Paragraph(project["title"], styles["Role"]))
        main_story.append(Paragraph(project["meta"], styles["Meta"]))
        for bullet in project["bullets"]:
            main_story.append(bullet_paragraph(bullet, styles))

    body_height = doc.height - header_height - 0.8 * cm
    sidebar_width = 5.4 * cm
    main_width = doc.width - sidebar_width

    body = Table(
        [
            [
                KeepInFrame(
                    sidebar_width - 14,
                    body_height - 12,
                    sidebar_story,
                    mode="shrink",
                ),
                KeepInFrame(
                    main_width - 18,
                    body_height - 12,
                    main_story,
                    mode="shrink",
                ),
            ]
        ],
        colWidths=[sidebar_width, main_width],
        rowHeights=[body_height],
    )
    body.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(SIDEBAR_BG)),
                ("BACKGROUND", (1, 0), (1, 0), colors.white),
                ("BOX", (0, 0), (-1, -1), 0, colors.white),
                ("LEFTPADDING", (0, 0), (0, 0), 14),
                ("RIGHTPADDING", (0, 0), (0, 0), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (1, 0), (1, 0), 16),
                ("RIGHTPADDING", (1, 0), (1, 0), 12),
            ]
        )
    )

    doc.build([header, Spacer(1, 0.18 * cm), body])


def main():
    build_docx()
    build_pdf()
    print(f"Created {DOCX_PATH}")
    print(f"Created {PDF_PATH}")


if __name__ == "__main__":
    main()
